from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from datetime import datetime
import os

def export_to_pdf(company_data, summary):
    """
    Export the valuation summary to a PDF file.
    """
    try:
        # Create output directory if it doesn't exist
        os.makedirs('output', exist_ok=True)
        
        # Sanitize company name for filename
        safe_name = "".join(c for c in company_data['name'] if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"output/valuation_summary_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph(f"Valuation Summary: {company_data['name']}", title_style))
        
        # Key Metrics
        metrics_style = ParagraphStyle(
            'Metrics',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=20
        )
        
        metrics = f"""
        Industry: {company_data['industry']}
        Revenue: ${company_data['revenue']:.2f}B
        Revenue Growth: {company_data['revenue_growth']:.1f}%
        Profit Margin: {company_data['profit_margin']:.1f}%
        """
        story.append(Paragraph("Key Metrics:", styles['Heading2']))
        story.append(Paragraph(metrics, metrics_style))
        
        # Summary
        story.append(Paragraph("Analysis:", styles['Heading2']))
        story.append(Paragraph(summary, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        return filename
        
    except Exception as e:
        raise Exception(f"Error creating PDF: {str(e)}")

def export_to_word(company_data, summary):
    """
    Export the valuation summary to a Word document.
    """
    try:
        # Create output directory if it doesn't exist
        os.makedirs('output', exist_ok=True)
        
        # Sanitize company name for filename
        safe_name = "".join(c for c in company_data['name'] if c.isalnum() or c in (' ', '-', '_')).strip()
        filename = f"output/valuation_summary_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        
        # Title
        doc.add_heading(f"Valuation Summary: {company_data['name']}", 0)
        
        # Key Metrics
        doc.add_heading("Key Metrics", level=1)
        metrics = [
            f"Industry: {company_data['industry']}",
            f"Revenue: ${company_data['revenue']:.2f}B",
            f"Revenue Growth: {company_data['revenue_growth']:.1f}%",
            f"Profit Margin: {company_data['profit_margin']:.1f}%"
        ]
        
        for metric in metrics:
            doc.add_paragraph(metric, style='List Bullet')
        
        # Summary
        doc.add_heading("Analysis", level=1)
        doc.add_paragraph(summary)
        
        # Save document
        doc.save(filename)
        return filename
        
    except Exception as e:
        raise Exception(f"Error creating Word document: {str(e)}")
